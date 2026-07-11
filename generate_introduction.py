from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Page margins
section = doc.sections[0]
section.page_height = Inches(11)
section.page_width  = Inches(8.5)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Heading style fixup helpers
def add_heading1(text):
    p = doc.add_heading(level=1)
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_heading2(text):
    p = doc.add_heading(level=2)
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_heading3(text):
    p = doc.add_heading(level=3)
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_para(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def bullet(label, body):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(label)
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(12); r1.bold = True
    if body:
        r2 = p.add_run(" " + body)
        r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
    return p

def numbered(label, body):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(label)
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(12); r1.bold = True
    if body:
        r2 = p.add_run(" " + body)
        r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
    return p

# ─────────────────────────────────────────────────────────────────────────────
# INTRODUCTION TEXT
# ─────────────────────────────────────────────────────────────────────────────
add_heading1('1. Introduction')

add_heading2('1.1. Background and Context')
add_para(
    "The global construction and building sector is responsible for a significant proportion of total global energy consumption and greenhouse gas (GHG) emissions. "
    "With rapid urbanization and economic growth, this trend is particularly pronounced in developing nations such as India. "
    "In India, the building sector accounts for over 30% of total electricity consumption, a figure that is projected to rise exponentially as the demand for cooling and modernized infrastructure increases (IEA, 2021). "
    "To mitigate the environmental impact of this growth, it is imperative to design energy-efficient buildings that minimize the Energy Use Intensity (EUI), measured in kWh/m²·yr. "
    "Achieving optimal EUI requires rigorous consideration of the building envelope, materials, and localized climate conditions during the early design phases."
)
add_para(
    "While conventional physics-based Building Energy Simulation (BES) tools (e.g., EnergyPlus, TRNSYS) offer high-fidelity performance analyses, they are computationally intensive, require highly detailed inputs that are often unavailable early in the design cycle, and require specialized domain expertise (Clarke, 2001). "
    "Consequently, there is an industry-wide reliance on post-design evaluation rather than proactive, prescriptive energy optimization. "
    "Furthermore, architects and builders in India face unique challenges due to diverse climate zones (ranging from composite and hot-dry to warm-humid) and a lack of highly localized, accessible databases for building material thermal properties and Life Cycle Cost Analysis (LCCA)."
)

add_heading2('1.2. Problem Statement')
add_para(
    "Despite advancements in sustainable architecture, several critical gaps remain in the current building design workflow:"
)
bullet("Computational Bottleneck in Early Design:", "Traditional dynamic simulations are too slow and complex to iterate over thousands of potential design configurations during the conceptual phase.")
bullet("Lack of Localized Context:", "Generic material libraries often fail to reflect regional building practices and standards, such as those prescribed by the Building Materials & Technology Promotion Council (BMTPC) and the Central Public Works Department (CPWD) of India.")
bullet("Disconnect Between Energy, Cost, and Carbon:", "Decision-makers often evaluate energy efficiency independently of capital costs, lifecycle financial implications (Net Present Value), and embodied carbon footprint, leading to suboptimal trade-offs.")
bullet("Black-Box Machine Learning Models:", "While data-driven surrogate models have emerged as faster alternatives to BES, their lack of interpretability prevents designers from understanding how specific architectural parameters (e.g., Window-to-Wall Ratio, insulation thickness) influence the predicted EUI.")

add_heading2('1.3. Proposed Solution')
add_para(
    "To address these systemic challenges, this research proposes a comprehensive, climate-aware Machine Learning (ML) framework tailored for the Indian context: the Climate-aware Material Recommendation & EUI Predictor. "
    "We present an end-to-end predictive and prescriptive AI engine that fundamentally transforms how architects approach early-stage design optimization."
)
add_para(
    "Our proposed system replaces computationally heavy physical simulations with a robust Extreme Gradient Boosting (XGBoost) regression surrogate model. "
    "Trained on a diverse synthetic dataset of physics-based simulations covering India's varied climate zones, this model predicts annual EUI near-instantaneously with high fidelity. "
    "To ensure environmental accuracy, the system auto-resolves geographic coordinates and integrates dynamically with the NASA POWER Climatology API to retrieve long-term weather markers (Cooling Degree Days, solar irradiance, and temperature profiles). "
    "Crucially, the framework is seeded with Indian material standards (BMTPC, CPWD/BIS) to ensure realistic structural assemblies and compliance checks against the Energy Conservation Building Code (ECBC) 2017."
)
add_para(
    "Beyond prediction, the system acts as an automated design consultant. "
    "It utilizes a Generative AI Optimizer that batches up to 5,000 design iterations through the ML surrogate in seconds, exploring the vast multi-dimensional design space to identify Pareto-optimal configurations that balance energy consumption, embodied carbon, and 30-year lifecycle costs."
)

add_heading2('1.4. Key Contributions')
add_para(
    "This research introduces the following core contributions to the fields of building science and applied artificial intelligence:"
)
bullet("Rapid EUI Prediction Surrogate:", "Development of a highly accurate XGBoost regressor that dramatically accelerates EUI forecasting without compromising the precision expected from physics-based engines (R² ~0.91).")
bullet("Multi-Objective Generative Optimization:", "Implementation of a sophisticated optimization engine that discovers Pareto-optimal material and envelope configurations, explicitly balancing energy performance against economic constraints (30-Year LCCA NPV) and environmental impact (embodied carbon).")
bullet("Explainable AI (XAI) Integration:", "Integration of SHAP (SHapley Additive exPlanations) to demystify the ML model, providing designers with granular, actionable insights into the per-prediction feature importance of their architectural choices.")
bullet("Localized Climatology and Material Standards:", "A novel fusion of real-time NASA POWER climate data with Indian regulatory material properties (BMTPC/CPWD), enabling automated ECBC 2017 compliance verification and context-specific recommendations.")

add_heading2('1.5. Document Organization')
add_para(
    "The remainder of this paper is structured as follows. Section 2 reviews the related work in building energy modeling, surrogate models, and multi-objective optimization. "
    "Section 3 details the methodology, encompassing the physics-informed synthetic data generation, the XGBoost ML pipeline, and the mathematical formulations for the generative optimizer and LCCA. "
    "Section 4 presents the experimental results, validation of the surrogate model, and a case study demonstrating the framework's application. "
    "Finally, Section 5 concludes the research with a summary of the impact and directions for future development."
)

add_heading1('References')
add_para("Clarke, J. A. (2001). Energy Simulation in Building Design. Butterworth-Heinemann.")
add_para("IEA (2021). India Energy Outlook 2021, IEA, Paris. https://www.iea.org/reports/india-energy-outlook-2021")
add_para("Bureau of Energy Efficiency (2017). Energy Conservation Building Code (ECBC). Ministry of Power, Government of India.")
add_para("BMTPC (2021). Thermal Properties of Building Materials Technical Document. Building Materials & Technology Promotion Council, India.")
add_para("Lundberg, S. M., & Lee, S.-I. (2017). A Unified Approach to Interpreting Model Predictions. Advances in Neural Information Processing Systems, 30.")
add_para("Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining.")

doc.save('Introduction_Section.docx')
print("Introduction section docx generated successfully.")
