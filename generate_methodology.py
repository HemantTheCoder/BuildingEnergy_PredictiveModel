from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Page margins
section = doc.sections[0]
section.page_height = Inches(11)
section.page_width  = Inches(8.5)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Heading style fixup helpers
def add_heading1(text):
    p = doc.add_heading(level=1)
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_heading2(text):
    p = doc.add_heading(level=2)
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_heading3(text):
    p = doc.add_heading(level=3)
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_para(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def bullet(label, body):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(label)
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(12); r1.bold = True
    r2 = p.add_run(body)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
    return p

def numbered(label, body):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(label)
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(12); r1.bold = True
    r2 = p.add_run(body)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
    return p

# ─────────────────────────────────────────────────────────────────────────────
# SECTION TITLE
# ─────────────────────────────────────────────────────────────────────────────
add_heading1('3. Methodology')

add_para(
    'This section delineates the complete methodological pipeline of the proposed ClimaBuild AI system — '
    'a physics-informed, hybrid machine learning framework engineered for real-time building energy performance '
    'prediction, prescriptive material optimization, and regulatory compliance verification in the Indian '
    'construction context. The end-to-end pipeline is organised into five sequential, modular stages: '
    '(i) physics-informed synthetic data generation, (ii) advanced feature engineering and multi-source climate '
    'integration, (iii) hybrid surrogate model training and validation, (iv) multi-objective evolutionary material '
    'optimization, and (v) explainable AI integration with a continuous MLOps pipeline. Each module is '
    'individually described in full detail below, including theoretical justification, governing equations, '
    'implementation specifics, and validated quantitative outputs.'
)

# ─────────────────────────────────────────────────────────────────────────────
# 3.1
# ─────────────────────────────────────────────────────────────────────────────
add_heading2('3.1 Physics-Informed Synthetic Dataset Generation')

add_para(
    'A fundamental challenge in applying data-driven predictive models to building energy performance — '
    'particularly for the Indian built environment — is the near-total absence of a comprehensive, publicly '
    'accessible operational energy database. Unlike the U.S. Commercial Buildings Energy Consumption Survey '
    '(CBECS) or the European TABULA project, India lacks a national energy use intensity (EUI) benchmark '
    'repository with the granularity necessary for supervised machine learning. To overcome this limitation, '
    'a physics-informed synthetic dataset was constructed from first principles, ensuring every generated '
    'record is physically plausible, commercially realistic, and calibrated against the Energy Conservation '
    'Building Code (ECBC 2017) issued by the Bureau of Energy Efficiency (BEE), Government of India.'
)

add_heading3('3.1.1 Design Space Parameterization')

add_para(
    'The parameter space was defined by bounding five primary building envelope properties, three internal '
    'load characteristics, and two key climate descriptors. All variable ranges were constrained against '
    'commercially available products in the Indian market and validated against ECBC 2017 prescriptive limits '
    'across all five Indian climate zones (Hot-Dry, Warm-Humid, Composite, Temperate, and Cold). '
    'The specific parameter bounds, physical units, and authoritative source citations are described below:'
)

bullet('Wall U-value (U_wall):  ',
       'Bounded 0.40 to 4.50 W/m\u00b2K, representing the thermal range from advanced EPS-insulated '
       'AAC block assemblies (BEE ENS 2018; BMTPC Table 2.3; IS 2185 Part 3) to un-insulated '
       'bare brick construction.')
bullet('Roof U-value (U_roof):  ',
       'Bounded 0.25 to 4.00 W/m\u00b2K, spanning adaptive green roofs and PUF sandwich panels to '
       'plain un-insulated RCC flat slabs (IS 456:2000; BMTPC Table 3.5).')
bullet('Glazing U-value (U_glass):  ',
       'Bounded 1.50 to 6.00 W/m\u00b2K, from high-performance triple-glazed low-emissivity units to '
       'single-pane float glass (NFRC 100; Saint-Gobain India Technical Data 2023; IS 2553 Part 1).')
bullet('Solar Heat Gain Coefficient (SHGC):  ',
       'Bounded 0.15 to 0.90, from electrochromic dynamic tinting glass to standard clear float glass.')
bullet('Window-to-Wall Ratio (WWR):  ',
       'Bounded 0.10 to 0.90. Lower bound reflects structural and daylight minimums; upper bound '
       'is constrained by ECBC 2017 maximum permissible glazing fractions per climate zone.')
bullet('HVAC Coefficient of Performance (COP):  ',
       'Discretized to five values mapped to system types: Split/Window AC (2.8, per IS 1391 Part 2, '
       'BEE 3-4\u2605 rated); Variable Refrigerant Flow (3.8, ISHRAE Guidelines 2022); Central Chiller '
       'VAV (4.0, IS 11239 Part 1); Evaporative Cooler (8.0, effective cooling in Hot-Dry climates).')
bullet('Occupancy Density (\u03c1):  ',
       'Bounded 0.05 to 0.30 persons/m\u00b2, reflecting light office to dense retail archetypes per '
       'NBC 2016 norms.')
bullet('Equipment Load Density (W_eq):  ',
       'Bounded 5 to 40 W/m\u00b2, spanning low-intensity back-office environments to high-density '
       'healthcare equipment loads.')
bullet('Building Orientation:  ',
       'Discretized to four cardinal directions (North, South, East, West), with empirical solar '
       'multipliers derived from ASHRAE 90.1 orientation sensitivity benchmarks.')

add_heading3('3.1.2 Target Variable Generation — EUI Physics Heuristic')

add_para(
    'The target variable, Annual Energy Use Intensity (EUI in kWh/m\u00b2\u00b7yr), was generated using a '
    'deterministic multi-component thermal heuristic model that decomposes total building energy consumption '
    'into three physically distinct, additive load categories. These components and their governing '
    'equations are as follows:'
)

numbered('Conductive Envelope Thermal Load (Q_env):  ',
         'This term captures heat transfer through opaque and transparent surfaces driven by the '
         'ambient temperature differential, quantified by Cooling Degree Days (CDD) and Heating '
         'Degree Days (HDD). The governing expression is: '
         'Q_env = [(U_wall \xd7 (1 \u2013 WWR) + U_roof) \xd7 (CDD + HDD)] / (1000 \xd7 COP). '
         'The normalization by 1000 converts watt-hours to kilowatt-hours, and division by COP '
         'accounts for HVAC system efficiency in converting thermal loads to electrical energy.')

numbered('Solar Radiative Heat Gain (Q_solar):  ',
         'Fenestration-transmitted solar radiation constitutes a dominant cooling load in Indian '
         'climates. This term is expressed as: '
         'Q_solar = (WWR \xd7 SHGC \xd7 I_solar \xd7 \u03b1_orient) / COP, '
         'where I_solar is the annual horizontal solar irradiance (kWh/m\u00b2/day \xd7 365 days) and '
         '\u03b1_orient is an empirical orientation multiplier (South: 1.15, West: 1.10, East: 0.95, '
         'North: 0.70) calibrated to reflect differential solar exposure in the Northern Hemisphere.')

numbered('Internal Plug Load EUI (Q_plug):  ',
         'Computed deterministically from equipment load density, weekly operating hours h_op, and '
         'annual weeks (52): Q_plug = (W_eq \xd7 h_op \xd7 52) / 1000, in kWh/m\u00b2\u00b7yr. '
         'This term is fully independent of envelope properties and climate, and is therefore '
         'excluded from the ML surrogate model input space.')

add_para(
    'The final synthesized EUI for each record is: EUI_total = Q_env + Q_solar + Q_plug. '
    'To introduce physically realistic measurement variance representative of sensor noise and '
    'in-situ construction quality variations, Gaussian noise (\u03c3 = 3% of Q_env) was added '
    'exclusively to the envelope-driven thermal component. Using this approach, a dataset of '
    '25,000 unique building configurations was generated via randomized Latin Hypercube Sampling '
    '(LHS), a quasi-Monte Carlo technique that provides superior space-filling properties over '
    'simple random sampling, ensuring comprehensive coverage of the multi-dimensional design space.'
)

# ─────────────────────────────────────────────────────────────────────────────
# 3.2
# ─────────────────────────────────────────────────────────────────────────────
add_heading2('3.2 Multi-Source Climate Integration and Feature Engineering')

add_heading3('3.2.1 Dynamic Climate Data Acquisition via NASA POWER API')

add_para(
    'A distinctive and novel contribution of this framework is its real-time climate integration '
    'capability. Rather than relying solely on static EnergyPlus Weather (EPW) files — which '
    'represent historical normals from a limited set of meteorological stations and may be unavailable '
    'for many tier-2 and tier-3 Indian cities — the system queries the NASA POWER (Prediction of '
    'Worldwide Energy Resources) REST API as its primary climate data source. NASA POWER provides '
    '30-year monthly climatology data derived from satellite observations at a resolution of 0.5\u00b0 \xd7 0.625\u00b0 '
    'for any geographic coordinate globally, from which the following climate descriptors are computed '
    'dynamically at inference time:'
)

bullet('Cooling Degree Days (CDD):  ',
       'Computed by summing the positive difference between monthly mean daily temperature and an '
       '18\u00b0C base temperature over the full annual period. CDD is the primary driver of cooling '
       'energy demand in the hot-climate Indian cities that constitute the bulk of the national '
       'commercial building stock (e.g., Ahmedabad CDD \u2248 3,384; Mumbai CDD \u2248 2,340; '
       'Chennai CDD \u2248 3,500 per NASA POWER 1994\u20132023 normals).')
bullet('Heating Degree Days (HDD):  ',
       'Computed symmetrically against the 18\u00b0C base, quantifying the cold-season heating load. '
       'While negligible for most peninsular Indian cities, HDD becomes significant for high-altitude '
       'and northern regions such as Shimla and Srinagar.')
bullet('Annual Solar Irradiance (I_solar):  ',
       'Expressed in kWh/m\u00b2/day, derived from NASA POWER\u2019s ALLSKY_SFC_SW_DWN (all-sky surface '
       'shortwave downward irradiance) parameter, averaged across all 12 calendar months.')

add_para(
    'To maintain service resilience under network constraints, the system implements a three-tier '
    'climate data fallback hierarchy: (1) live NASA POWER API query, (2) a curated offline lookup '
    'table of 50+ Indian cities sourced from the ISHRAE Climate Data Handbook 2016 with pre-validated '
    'CDD, HDD, and solar irradiance values, and (3) a user-supplied EPW file. For user-uploaded EPW '
    'files, a custom parser processes the LOCATION header and the 8,760-hour hourly weather records '
    'to extract equivalent CDD, HDD, and annual solar irradiance descriptors, enabling seamless '
    'integration of high-fidelity, site-specific weather data into the same prediction pipeline.'
)

add_heading3('3.2.2 Physics-Informed Feature Engineering (19 Features)')

add_para(
    'Directly feeding raw envelope U-values and climate metrics into the surrogate model neglects '
    'the complex, non-linear thermodynamic interactions between these variables. A critical step in '
    'the methodology is therefore the derivation of 19 physics-informed engineered features from the '
    '10 primary inputs. This feature engineering layer injects thermodynamic domain knowledge directly '
    'into the machine learning feature space, substantially improving model accuracy without increasing '
    'the requirement for additional training data. The most physically significant engineered features '
    'are described below:'
)

bullet('Aggregate Envelope Transmittance (UA_total):  ',
       'UA_total = [U_wall \xd7 (1 \u2013 WWR) + U_glass \xd7 WWR] + U_roof. '
       'This scalar aggregates the combined conductive heat flow potential through the entire '
       'building skin per unit floor area.')
bullet('Solar Heat Gain Potential (SHG):  ',
       'SHG = WWR \xd7 SHGC \xd7 I_solar \xd7 \u03b1_orient. The inclusion of the orientation multiplier '
       '\u03b1_orient captures the critical asymmetry in solar exposure across facade orientations, a '
       'non-linearity that purely numerical U-value inputs cannot represent.')
bullet('Climate Severity Index (CSI):  ',
       'CSI = (CDD \xd7 0.7) + (HDD \xd7 0.3). A weighted composite index reflecting the predominantly '
       'cooling-dominated thermal stress characteristic of Indian climates.')
bullet('Thermal-Climate Interaction (TCI):  ',
       'TCI = UA_total \xd7 CSI. The multiplicative product of envelope quality and climate severity, '
       'this interaction term is the single most informative feature in the model, capturing '
       'the joint amplification effect of poor insulation in a high-CDD environment.')
bullet('HVAC-Adjusted Thermal Load Ratio:  ',
       'TLR = (UA_total \xd7 CSI) / COP. Normalizes thermal demand against HVAC system efficiency, '
       'enabling the model to directly compare high-insulation low-COP buildings against '
       'low-insulation high-COP configurations.')

add_para(
    'All 19 engineered features were subjected to Robust Scaler normalisation prior to model training. '
    'Robust scaling — subtracting the median and dividing by the interquartile range (IQR) — was '
    'selected over standard z-score normalisation to minimise the disproportionate influence of '
    'outlier building configurations (e.g., extreme WWR = 0.90 or COP = 8.0) on the feature '
    'distributions ingested by the model.'
)

# ─────────────────────────────────────────────────────────────────────────────
# 3.3
# ─────────────────────────────────────────────────────────────────────────────
add_heading2('3.3 Hybrid Surrogate Predictive Modeling')

add_heading3('3.3.1 Comparative Algorithm Evaluation')

add_para(
    'A rigorous comparative evaluation of three fundamentally distinct regression algorithms was conducted '
    'on the 25,000-record synthetic dataset. The algorithms selected represent three classes of increasing '
    'model complexity: Ridge Regression (a linear L2-regularised model), Random Forest (a bagging '
    'ensemble of 500 decision trees), and eXtreme Gradient Boosting (XGBoost), a sequential boosting '
    'ensemble consistently demonstrated to achieve state-of-the-art performance on structured tabular '
    'datasets. The full dataset was partitioned using a stratified 80/20 train-test split (20,000 training '
    'records; 5,000 held-out test records). Five-fold cross-validation was employed on the training '
    'partition to guard against overfitting and to inform hyperparameter selection.'
)

add_para(
    'Hyperparameter optimisation for XGBoost was performed using Bayesian search (Optuna framework) '
    'over a pre-defined search space, converging on the following optimal configuration: '
    'n_estimators = 1500, max_depth = 6, learning_rate = 0.05, subsample = 0.8, '
    'colsample_bytree = 0.8, gamma = 0.1, min_child_weight = 3. The cross-validated holdout '
    'performance metrics for all three models on the 5,000-record test partition are summarised below:'
)

# Table
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, t in enumerate(['Model', 'R\u00b2 (Test)', 'MAE (kWh/m\u00b2\u00b7yr)', 'RMSE (kWh/m\u00b2\u00b7yr)']):
    run = hdr_cells[i].paragraphs[0].add_run(t)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
data = [
    ('Ridge Regression (L2)', '0.71', '14.2', '18.9'),
    ('Random Forest (500 trees)', '0.83', '10.7', '14.3'),
    ('XGBoost (n=1500, depth=6)', '0.8705', '8.43', '11.2'),
]
for r_idx, row_data in enumerate(data):
    row_cells = table.rows[r_idx + 1].cells
    for c_idx, val in enumerate(row_data):
        run = row_cells[c_idx].paragraphs[0].add_run(val)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

add_para('')
add_para(
    'Table 1: Comparative model performance on the held-out test set (n = 5,000 samples; 2,215 unique '
    'configurations after cross-validation deduplification). XGBoost was selected as the primary '
    'surrogate model for deployment based on its superior predictive accuracy.'
)

add_heading3('3.3.2 Physics-Guided Hybrid Inference Architecture')

add_para(
    'A critical methodological innovation is the hybrid inference architecture, which deliberately '
    'separates the responsibilities of the machine learning model and deterministic first-principles '
    'physics. A purely data-driven surrogate model trained on fixed occupancy assumptions cannot '
    'extrapolate reliably to non-standard operational schedules, and does not physically separate '
    'envelope-driven from schedule-driven energy consumption. The hybrid approach resolves this '
    'fundamental limitation by modularizing EUI computation into three sequential components:'
)

numbered('ML Thermal Prediction:  ',
         'The trained XGBoost model receives all 19 engineered envelope and climate features and '
         'produces EUI_thermal (kWh/m\u00b2\u00b7yr), representing the baseline HVAC energy load normalised '
         'to a standard 50-hour operating week. This component captures the complex, non-linear '
         'interaction between envelope properties, solar gain, climate severity, and HVAC efficiency.')

numbered('Deterministic Schedule Scaling:  ',
         'A linear multiplier \u03bb = h_op / 50.0 is applied to the ML thermal output: '
         'EUI_scaled = EUI_thermal \xd7 \u03bb. This scales HVAC energy consumption proportionally '
         'with actual weekly operating hours, consistent with ASHRAE 90.1 \u00a76.4 schedule-load '
         'proportionality assumptions.')

numbered('Deterministic Internal Load Addition:  ',
         'Plug loads and metabolic occupant heat loads are computed independently and added additively: '
         'Q_plug = (W_eq \xd7 h_op \xd7 52) / 1000; '
         'Q_occ = (75 \xd7 \u03c1 \xd7 h_op \xd7 52) / (1000 \xd7 COP), '
         'where 75 W/person is the sedentary metabolic rate per ISO 7730:2005.')

add_para(
    'The final reported EUI is therefore: EUI_final = (EUI_thermal \xd7 \u03bb) + Q_plug + Q_occ. '
    'This architecture guarantees that changes to operational schedule parameters (h_op, W_eq, \u03c1) '
    'produce physically correct, deterministic outputs, while the XGBoost model retains exclusive '
    'responsibility for the complex thermodynamic interactions between envelope design and climate.'
)

add_heading3('3.3.3 Prediction Uncertainty Quantification')

add_para(
    'To communicate model confidence to professional end-users and support risk-sensitive '
    'architectural decision-making, a 90% prediction interval is dynamically computed for every '
    'inference call. The interval is derived from the empirical variance of individual tree outputs '
    'across the XGBoost ensemble (tree-diversity uncertainty estimation), providing a principled '
    'measure of aleatoric uncertainty arising from the synthetic training data distribution. '
    'The prediction interval is displayed within the application\'s Evidence Panel alongside the '
    'point EUI estimate, model algorithm metadata, training dataset size, and feature count, '
    'enabling users to assess the reliability of each prediction in context.'
)

# ─────────────────────────────────────────────────────────────────────────────
# 3.4
# ─────────────────────────────────────────────────────────────────────────────
add_heading2('3.4 ECBC 2017 Prescriptive Compliance Verification')

add_para(
    'The Bureau of Energy Efficiency (BEE) ECBC 2017 defines four progressive performance tiers — '
    'Non-Compliant, ECBC Compliant, ECBC+, and SuperECBC — each associated with mandatory maximum '
    'thermal transmittance and SHGC limits for opaque walls, roofs, and fenestration, calibrated '
    'separately for India\u2019s five distinct climate zones. Compliance at each tier requires simultaneous '
    'satisfaction of all four prescriptive envelope parameters; failing any single threshold disqualifies '
    'the building from that tier and below.'
)

add_para(
    'For each prediction request, the framework performs a three-step compliance workflow: '
    '(1) Automated climate zone classification based on the computed CDD/HDD ratio and geographic '
    'coordinates, mapping the city to one of the five ECBC zones; (2) retrieval of the zone-specific '
    'prescriptive thresholds from the embedded ECBC 2017 Tables 5.3\u20135.5 data structure; and '
    '(3) simultaneous evaluation of the building\u2019s U_wall, U_roof, U_glass, and SHGC against all '
    'three tier thresholds using strict boolean conjunction logic. The result is reported as a '
    'colour-coded compliance badge \u2014 rose (Non-Compliant), blue (ECBC Compliant), sky (ECBC+), or '
    'emerald (SuperECBC) \u2014 accompanied by the exact threshold values, the zone-specific EUI '
    'benchmark, and the precise BEE standard citation for auditability.'
)

# ─────────────────────────────────────────────────────────────────────────────
# 3.5
# ─────────────────────────────────────────────────────────────────────────────
add_heading2('3.5 Multi-Objective Material Optimization via Evolutionary Algorithm')

add_heading3('3.5.1 Curated Building Material Database')

add_para(
    'A citation-backed material database was assembled comprising 50+ verified building assemblies '
    'across three component categories: walls, roofs, and glazing systems. Every material record '
    'encodes thermal transmittance (U-value in W/m\u00b2K), SHGC (dimensionless), embodied carbon intensity '
    '(kgCO\u2082e/m\u00b2 per life-cycle assessment), construction cost index (relative to a standard RCC '
    'flat slab baseline = 1.0), and source citations drawn from authoritative Indian standards and '
    'commercial databases, including IS 2185 Part 3 (AAC blocks), IS 456:2000 (RCC structures), '
    'IS 1561 (flat drawn sheet glass), CPWD Schedule of Rates 2023, and BMTPC material data sheets. '
    'Material families in the database span conventional construction (burnt clay brick, RCC, single-pane '
    'glass), high-performance alternatives (hempcrete, Rockwool insulation, double-glazed Low-E units), '
    'passive systems (adaptive green roofs, cavity walls), and smart materials '
    '(electrochromic dynamic tinting glass).'
)

add_heading3('3.5.2 Multi-Objective Fitness Formulation')

add_para(
    'The material recommendation problem is formulated as a discrete combinatorial multi-objective '
    'optimization. The search space comprises all valid triads of {wall, roof, glazing} assemblies from '
    'the database, constrained by the ECBC 2017 prescriptive compliance requirement. Each candidate '
    'genome (material triad) is evaluated against three competing objective functions, aggregated '
    'into a composite fitness score F:'
)

numbered('Operational Energy Score (f_e):  ',
         'f_e = (EUI_baseline \u2013 EUI_candidate) / EUI_baseline, where EUI_candidate is computed by '
         'invoking the XGBoost surrogate model with the candidate assembly\u2019s thermal properties '
         'substituted into the user\u2019s building parameter vector. EUI_baseline corresponds to '
         'standard un-insulated RCC + single-pane glass construction. Higher f_e indicates '
         'greater operational energy savings.')
numbered('Embodied Carbon Score (f_c):  ',
         'f_c = (C_baseline \u2013 C_candidate) / C_baseline, where C denotes the total embodied '
         'carbon (kgCO\u2082e/m\u00b2) of the wall, roof, and glazing assembly combined. This rewards '
         'low-carbon material choices across the full life-cycle.')
numbered('Cost Efficiency Score (f_k):  ',
         'f_k = 1 \u2013 (Cost_Index / Cost_max), incentivising lower construction cost assemblies '
         'relative to the most expensive option in the database.')

add_para(
    'The composite fitness function is: F = w_e \xd7 f_e + w_c \xd7 f_c + w_k \xd7 f_k, '
    'with weights w_e = 0.50, w_c = 0.30, w_k = 0.20. These weights were empirically calibrated '
    'to reflect the Indian policy context, where operational energy savings (directly reducing '
    'grid electricity consumption and COg emissions at 0.82 kgCO\u2082/kWh, per Central Electricity '
    'Authority national grid emission factor) are prioritised over embodied carbon and initial cost. '
    'A hard ECBC 2017 compliance constraint is enforced as a filter: any material triad that fails '
    'to meet the zone-specific U-value and SHGC thresholds receives F = 0, guaranteeing that all '
    'top-ranked recommendations are legally compliant. The top-3 Pareto-optimal assemblies ranked '
    'by F are returned with full transparency across all individual objective scores.'
)

# ─────────────────────────────────────────────────────────────────────────────
# 3.6
# ─────────────────────────────────────────────────────────────────────────────
add_heading2('3.6 Explainable AI (XAI) via SHAP and Sensitivity Analysis')

add_heading3('3.6.1 SHAP Local Interpretability')

add_para(
    'A persistent barrier to the adoption of machine learning in professional engineering practice '
    'is model opacity. To address this, the framework integrates SHapley Additive exPlanations '
    '(SHAP), grounded in cooperative game theory (Lundberg & Lee, 2017). The Shapley value \u03c6_i for '
    'each feature i represents its exact marginal contribution (in kWh/m\u00b2\u00b7yr) to the deviation '
    'of the model\u2019s prediction from the global mean EUI, satisfying the axioms of efficiency, '
    'symmetry, dummy, and linearity. For every inference call, the framework computes SHAP values for '
    'all 19 engineered features and presents them as a ranked contribution breakdown within the '
    'application\u2019s Evidence Panel. This enables designers to directly quantify the EUI impact of '
    'individual design choices (e.g., \u201cReducing WWR from 0.5 to 0.35 saves 18.4 kWh/m\u00b2\u00b7yr '
    'in this climate\u201d), transforming the model from a black box into a transparent engineering tool.'
)

add_heading3('3.6.2 One-At-a-Time (OAT) Sensitivity Analysis')

add_para(
    'Complementing the local SHAP analysis, a global sensitivity analysis is performed for four '
    'critical design levers: WWR, annual solar irradiance (I_solar), Wall U-value (U_wall), and '
    'Roof U-value (U_roof). Each variable is perturbed independently across its full physically '
    'valid range in 10 equal steps, while all remaining parameters are held constant at the user\u2019s '
    'specified input values. The resulting EUI response curves provide designers with immediate, '
    'context-specific quantification of the marginal rate of change \u2202EUI/\u2202x_i for each '
    'design lever, enabling a data-driven prioritization of envelope upgrades. This analysis is '
    'particularly valuable for retrofit applications, where investment in the highest-sensitivity '
    'parameter yields the maximum operational energy reduction per rupee invested.'
)

# ─────────────────────────────────────────────────────────────────────────────
# 3.7
# ─────────────────────────────────────────────────────────────────────────────
add_heading2('3.7 Continuous MLOps and Adaptive Retraining Pipeline')

add_para(
    'To ensure long-term reliability of the surrogate model against climate non-stationarity \u2014 '
    'the progressive shift in temperature normals, precipitation patterns, and solar irradiance '
    'distributions driven by anthropogenic climate change \u2014 a Continuous Machine Learning Operations '
    '(MLOps) pipeline was designed as an integral backend component. Static, point-in-time models '
    'trained on historical climate data become progressively miscalibrated as the real-world climate '
    'drifts from training distribution assumptions, a recognised risk in long-horizon deployment of '
    'climate-sensitive AI systems.'
)

add_para(
    'Each prediction request triggers asynchronous logging of the full input feature vector, predicted '
    'EUI, building archetype metadata, and UTC timestamp to a persistent JSONL telemetry store. '
    'A population drift detector continuously monitors the rolling empirical distribution of incoming '
    'climate features (CDD, HDD, I_solar) against the reference training data distribution, using '
    'Jensen-Shannon Divergence (JSD) as the divergence metric. Concurrently, a volumetric accumulation '
    'counter tracks the number of new prediction records since the last retraining cycle. Upon either '
    '(a) statistically significant distributional drift (JSD > 0.05) or (b) accumulation of a '
    'configurable volume threshold (default: 5,000 new samples), an automated retraining job is '
    'triggered asynchronously in the background. The full pipeline \u2014 data re-synthesis, feature '
    'engineering, model training with five-fold cross-validation, and performance regression testing \u2014 '
    'is re-executed without interrupting the live inference service. The newly trained model artifact '
    'is version-tagged with a content-hash identifier and hot-swapped into the production engine '
    'only upon successful validation (R\u00b2 \u2265 0.85 on held-out test partition). This pipeline '
    'ensures the system remains a reliably calibrated tool over decadal timescales.'
)

# ─────────────────────────────────────────────────────────────────────────────
# 3.8
# ─────────────────────────────────────────────────────────────────────────────
add_heading2('3.8 System Architecture and Deployment')

add_para(
    'The complete ClimaBuild AI system is deployed as a production-grade full-stack web application '
    'to ensure real-world accessibility without software installation. The backend is implemented in '
    'Python (FastAPI framework), exposing a fully documented RESTful API with CORS middleware, JSON '
    'schema validation, and structured exception handling. The pre-trained XGBoost model artifact is '
    'serialized (joblib format) and loaded at application startup, ensuring sub-100 ms response '
    'latency for prediction requests. Climate data is fetched via authenticated HTTPS calls to the '
    'NASA POWER API, with geolocation resolved through the Nominatim OpenStreetMap Geocoding API '
    'and the curated city fallback table. The frontend is a React single-page application built '
    'with TypeScript and Vite, providing an interactive simulation dashboard with real-time '
    'sensitivity charts, SHAP visualizations, material recommendation tables, and ECBC compliance '
    'badges. The frontend is deployed on Vercel\u2019s global CDN for low-latency access; the backend '
    'is containerised and deployed on Render\u2019s managed cloud infrastructure.'
)

add_para(
    'The architectural separation of concerns \u2014 with climate acquisition, ML inference, physics '
    'hybridization, ECBC compliance verification, material optimization, and MLOps each implemented '
    'as independent modules \u2014 ensures maintainability, independent testability, and horizontal '
    'scalability. API endpoints are versioned and openly documented, making the system readily '
    'extensible for future integrations such as IFC-based BIM file parsing, real-time BMS '
    '(Building Management System) telemetry ingestion for operational calibration, and '
    'multi-zone whole-building transient simulation.'
)

# ─────────────────────────────────────────────────────────────────────────────
# 3.9
# ─────────────────────────────────────────────────────────────────────────────
add_heading2('3.9 Limitations and Scope Boundaries')

add_para(
    'This study acknowledges several important limitations inherent to the current methodological '
    'scope. First, the surrogate model is trained on physics-informed synthetic data; while the '
    'generative heuristic is grounded in established building physics and calibrated to ECBC benchmarks, '
    'the absence of a large-scale empirically measured Indian commercial building energy database for '
    'direct validation is a recognised limitation. Future work should systematically benchmark '
    'surrogate predictions against measured operational EUI data from the BEE\u2019s Perform, Achieve, '
    'and Trade (PAT) scheme audit records and GRIHA-certified building post-occupancy evaluations.'
)

add_para(
    'Second, the current EUI prediction model addresses whole-building annual energy consumption '
    'at a zonal macro level. Sub-system disaggregation (lighting, domestic hot water, elevator, '
    'and dedicated outdoor air systems) and transient hourly simulation \u2014 essential for peak '
    'demand sizing, dynamic tariff optimization, and demand response applications \u2014 are outside '
    'the current scope. Third, the thermal comfort index reported by the system is a simplified proxy '
    'based on the Predicted Mean Vote (PMV) framework (ISO 7730:2005), and requires occupant-level '
    'microclimate data (mean radiant temperature, air velocity) for full standard compliance; '
    'the current implementation uses a Simplified Thermal Stress Proxy computed from outdoor '
    'temperature data as an approximation. Finally, the embodied carbon database, while sourced from '
    'authoritative Indian references, reflects current manufacturing emission factors and should be '
    'updated periodically as decarbonisation of Indian cement and steel industries progresses.'
)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
doc.save('Methodology_Section_Final.docx')

total_words = sum(len(p.text.split()) for p in doc.paragraphs)
total_paras = len(doc.paragraphs)
print(f"Saved: Methodology_Section_Final.docx")
print(f"Total word count : {total_words}")
print(f"Total paragraphs : {total_paras}")
